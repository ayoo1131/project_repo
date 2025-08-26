#fill_template.py
import os
from docx import Document
import io
from pathlib import Path
import logging

def fill_template(username, cover_letter_input):
    current_file_path = os.path.abspath(__file__)
    parent_dir_path = os.path.dirname(os.path.dirname(current_file_path)) #Full path to /backend directory
    template_doc_path = None #create variable for filepath to coverletter template. 
    
    if (cover_letter_input['social'] != '' and cover_letter_input['extra'] != ''): #extra and social fields are NOT empty
        template_doc_path = parent_dir_path + '/templates/cover_letter_template_all.docx'
    
    elif (cover_letter_input['social'] == '' and cover_letter_input['extra'] == ''): #extra and social fields are empty
        template_doc_path = parent_dir_path + '/templates/cover_letter_template_no_social_extra.docx'

    elif (cover_letter_input['social'] != '' and cover_letter_input['extra'] == ''): #extra field is empty
        template_doc_path = parent_dir_path + '/templates/cover_letter_template_no_extra.docx'

    elif (cover_letter_input['social'] == '' and cover_letter_input['extra'] != ''): #social field is empty
        template_doc_path = parent_dir_path + '/templates/cover_letter_template_no_social.docx'

    doc = Document(template_doc_path) #Document() creates or opens existing Word document. This opens preexisting document
    
    for para in doc.paragraphs: #Loop over all blocks of text seperated by line break. 
        for run in para.runs: #Loop over all different font, size, formatting in paragraph. Usually only 1.
            for input_name, input_value in cover_letter_input.items(): #loop through key and value in cover_letter_input dictionary, check if run has <> block 
                if (f'<{input_name}>') in run.text:
                    run.text = run.text.replace(f'<{input_name}>', input_value)

    #This will save document to disk, not necessary for user download
    #doc.save(filePath) #Saves the document to the disk at the indicated filePath
    
    buffer = io.BytesIO() #Creates an in-memory buffer(like virtual file) to store binary data
    doc.save(buffer) #Saves the same document into in-memory buffer. Used to send document to user by HTTP response w/o touching file system/disk.
    buffer.seek(0) #Above line moves file pointer to end of buffer, current line moves pointer to beginning of in-memory buffer. Technically not necessary, but good practice to be able to read document again.

    return buffer
